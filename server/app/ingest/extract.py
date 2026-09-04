"""Bytes -> `RawDocument`. The only module that touches file formats.

Format is sniffed from magic bytes, never from the extension: users rename
things, browsers mislabel uploads, and "resume.pdf" is a .docx often enough to
matter. A wrong guess here produces a `CORRUPT` warning instead of silently
extracting nothing.

Nothing in this module interprets the text. It reports what it read and how
confident it is that it read it correctly.
"""

from __future__ import annotations

from ..schemas.ingest import (
    MIN_USABLE_CHARS,
    TYPICAL_MAX_PAGES,
    DocumentKind,
    ExtractionWarning,
    Layout,
    RawDocument,
    WarningCode,
    doc_id_for,
    sha256_bytes,
)
from .columns import TextLine, read_lines
from .normalize import garble_ratio, normalize

#: Past this share of unmapped glyphs the embedded font is broken and the text,
#: though non-empty, does not say what the page says.
GARBLE_BLOCK_RATIO = 0.05

_PDF_MAGIC = b"%PDF"
_ZIP_MAGIC = b"PK\x03\x04"


def sniff(data: bytes, filename: str | None = None) -> DocumentKind | None:
    """Detect the format from content. None means unsupported."""
    head = data[:8]
    if head.startswith(_PDF_MAGIC):
        return DocumentKind.PDF
    if head.startswith(_ZIP_MAGIC):
        # Every OOXML file is a zip; only a .docx has the main document part.
        # .pptx and .xlsx would otherwise reach the docx reader and raise.
        import zipfile
        from io import BytesIO

        try:
            with zipfile.ZipFile(BytesIO(data)) as z:
                if "word/document.xml" in z.namelist():
                    return DocumentKind.DOCX
        except zipfile.BadZipFile:
            return None
        return None
    try:
        data.decode("utf-8")
    except UnicodeDecodeError:
        return None
    return DocumentKind.TEXT


def extract(data: bytes, filename: str | None = None) -> RawDocument:
    """Extract text from an uploaded file, with warnings attached."""
    digest = sha256_bytes(data)
    kind = sniff(data, filename)

    if kind is None:
        return RawDocument(
            doc_id=doc_id_for(digest),
            kind=DocumentKind.TEXT,
            filename=filename,
            text="",
            sha256=digest,
            warnings=[
                ExtractionWarning(
                    code=WarningCode.UNSUPPORTED_TYPE,
                    message=(
                        "That file isn't a PDF, Word document, or plain text. "
                        "Export your résumé as PDF, or paste the text instead."
                    ),
                    blocking=True,
                )
            ],
        )

    if kind is DocumentKind.PDF:
        text, pages, layout, warnings = _read_pdf(data)
    elif kind is DocumentKind.DOCX:
        text, pages, layout, warnings = _read_docx(data)
    else:
        text, pages, layout, warnings = data.decode("utf-8", errors="replace"), 0, Layout.UNKNOWN, []

    text = normalize(text)
    warnings.extend(_audit(text, kind, pages))

    return RawDocument(
        doc_id=doc_id_for(digest),
        kind=kind,
        filename=filename,
        text=text,
        page_count=pages,
        layout=layout,
        warnings=warnings,
        sha256=digest,
    )


def extract_pasted(text: str, filename: str | None = None) -> RawDocument:
    """The escape hatch. No extraction, so no layout risk and no font risk.

    Worth keeping first-class: when a PDF is a scan or an exotic template, asking
    the user to paste is a better answer than shipping an OCR dependency.
    """
    raw = text.encode("utf-8")
    digest = sha256_bytes(raw)
    clean = normalize(text)
    return RawDocument(
        doc_id=doc_id_for(digest),
        kind=DocumentKind.PASTED,
        filename=filename,
        text=clean,
        page_count=0,
        layout=Layout.UNKNOWN,
        warnings=_audit(clean, DocumentKind.PASTED, 0),
        sha256=digest,
    )


# ── format readers ─────────────────────────────────────────────────────────────


def _read_pdf(data: bytes) -> tuple[str, int, Layout, list[ExtractionWarning]]:
    import pymupdf

    warnings: list[ExtractionWarning] = []
    try:
        doc = pymupdf.open(stream=data, filetype="pdf")
    except Exception as exc:  # noqa: BLE001 — any parse failure is the same to the user
        return (
            "",
            0,
            Layout.UNKNOWN,
            [
                ExtractionWarning(
                    code=WarningCode.CORRUPT,
                    message=f"That PDF could not be opened ({type(exc).__name__}). "
                    "Try re-exporting it, or paste the text instead.",
                    blocking=True,
                )
            ],
        )

    with doc:
        if doc.needs_pass:
            return (
                "",
                doc.page_count,
                Layout.UNKNOWN,
                [
                    ExtractionWarning(
                        code=WarningCode.ENCRYPTED,
                        message="That PDF is password-protected. Remove the password, "
                        "or paste the text instead.",
                        blocking=True,
                    )
                ],
            )

        pages: list[str] = []
        multi_column = False
        for page in doc:
            text, found = read_lines(_page_lines(page), page.rect.width)
            multi_column = multi_column or found
            pages.append(text)
        page_count = doc.page_count

    if multi_column:
        warnings.append(_multi_column_warning())

    text = "\n\n".join(pages)
    if multi_column:
        layout = Layout.MULTI_COLUMN
    elif text.strip():
        layout = Layout.SINGLE_COLUMN
    else:
        # No text means no layout was observed. Reporting "single_column" for a
        # scan would be a claim we never checked.
        layout = Layout.UNKNOWN

    # A page break is a paragraph break, never a sentence break: joining pages
    # with "\n" would glue the last bullet of page 1 to the first of page 2.
    return text, page_count, layout, warnings


def _multi_column_warning() -> ExtractionWarning:
    """Advisory, not blocking. We reordered the columns; we did not verify them."""
    return ExtractionWarning(
        code=WarningCode.MULTI_COLUMN,
        message=(
            "This résumé uses columns. We read each column separately, but please check "
            "the extracted text before confirming — column layouts are the most common "
            "source of mixed-up dates."
        ),
    )


def _page_lines(page) -> list[TextLine]:  # noqa: ANN001 — pymupdf exports no page type
    """Flatten one page to positioned lines.

    Lines, not blocks: MuPDF merges same-baseline text from both columns into one
    block, so a block bbox spans the gutter and hides the layout. See columns.py.
    """
    out: list[TextLine] = []
    for block in page.get_text("dict").get("blocks", []):
        # type 1 is an image; it has no "lines" and nothing to read.
        for line in block.get("lines", []):
            text = "".join(span.get("text", "") for span in line.get("spans", []))
            if not text.strip():
                continue
            x0, y0, x1, y1 = line["bbox"]
            out.append(TextLine(x0=x0, y0=y0, x1=x1, y1=y1, text=text))
    return out


def _read_docx(data: bytes) -> tuple[str, int, Layout, list[ExtractionWarning]]:
    from io import BytesIO

    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    try:
        document = docx.Document(BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        return (
            "",
            0,
            Layout.UNKNOWN,
            [
                ExtractionWarning(
                    code=WarningCode.CORRUPT,
                    message=f"That Word document could not be read ({type(exc).__name__}). "
                    "Save it as PDF, or paste the text instead.",
                    blocking=True,
                )
            ],
        )

    parts: list[str] = []

    # Contact details often live in the page header, which `document.paragraphs`
    # does not include — losing the email is worse than a duplicate line.
    for section in document.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        break

    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = Paragraph(child, document).text.strip()
            if text:
                parts.append(text)
        elif tag == "tbl":
            # Résumé templates use tables for two-column layouts, and
            # `document.paragraphs` skips table content entirely — the single
            # biggest cause of a .docx extracting as half a résumé.
            parts.append(_read_table(Table(child, document)))

    text = "\n".join(p for p in parts if p.strip())
    if not _has_layout_table(document):
        return text, 0, Layout.SINGLE_COLUMN, []

    # A table-based two-column layout is the most common way a .docx scrambles,
    # and unlike the PDF path we have no geometry to recover from — so the warning
    # is the whole mitigation. See `_read_table` for why we still read row-major.
    return text, 0, Layout.MULTI_COLUMN, [_multi_column_warning()]


def _read_table(table: Table) -> str:
    """Row-major, one cell per line.

    Cells go on their own lines rather than joined with a separator: a layout
    table's "row" is not a record, and " | " between a skills sidebar and a job
    bullet invents a relationship between them.

    Row-major even though a two-column *layout* table would read better
    column-major, because the two cases are not distinguishable and the common
    one is row-major:

      * A real two-column résumé template is almost always ONE row with two
        cells, each holding many paragraphs — for which row-major and
        column-major are identical.
      * A many-row two-column table is usually real data ("2019 - 2021 | Senior
        Engineer, Acme"), where column-major would tear every date away from its
        job. That is the same failure the PDF path exists to prevent.

    So row-major is right for both shapes that actually occur, and the layout
    warning covers the residual case.
    """
    lines: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text and (not lines or lines[-1] != text):  # merged cells repeat
                lines.append(text)
    return "\n".join(lines)


def _has_layout_table(document) -> bool:  # noqa: ANN001 — docx has no public type export
    """A table with 2+ columns and no header styling is almost always layout."""
    return any(len(t.columns) >= 2 for t in document.tables)


# ── quality audit ──────────────────────────────────────────────────────────────


def _audit(text: str, kind: DocumentKind, pages: int) -> list[ExtractionWarning]:
    out: list[ExtractionWarning] = []

    if not text.strip():
        out.append(
            ExtractionWarning(
                code=WarningCode.NO_TEXT_LAYER,
                message=(
                    "No text could be read from that file. If it's a scan or a photo, "
                    "the words are pixels, not characters — paste the text instead."
                )
                if kind is DocumentKind.PDF
                else "That file is empty.",
                blocking=True,
            )
        )
        return out

    if len(text) < MIN_USABLE_CHARS:
        out.append(
            ExtractionWarning(
                code=WarningCode.TOO_SHORT,
                message=(
                    f"Only {len(text)} characters were read — too little to be a résumé. "
                    "If the file is a scan, paste the text instead."
                ),
                blocking=True,
            )
        )

    ratio = garble_ratio(text)
    if ratio > GARBLE_BLOCK_RATIO:
        out.append(
            ExtractionWarning(
                code=WarningCode.GARBLED,
                message=(
                    f"{ratio:.0%} of the extracted text is unreadable, which usually means "
                    "the PDF's fonts aren't embedded properly. Re-export it, or paste the text."
                ),
                blocking=True,
            )
        )

    if pages > TYPICAL_MAX_PAGES:
        out.append(
            ExtractionWarning(
                code=WarningCode.MANY_PAGES,
                message=(
                    f"{pages} pages is long for a résumé — check you uploaded the right file. "
                    "We'll still read all of it."
                ),
            )
        )

    return out


__all__ = ["GARBLE_BLOCK_RATIO", "sniff", "extract", "extract_pasted"]
